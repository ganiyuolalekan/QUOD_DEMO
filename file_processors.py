import io
import os
from typing import Union, BinaryIO
from pathlib import Path
import tempfile
import streamlit as st
from enum import Enum

# Import libraries for different file types
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
    from docx import Document
except ImportError:
    docx = None
    Document = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    import chardet
except ImportError:
    chardet = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

# Try to import python-docx2txt for .doc files
try:
    import docx2txt
except ImportError:
    docx2txt = None


class SupportedFileTypes(Enum):
    """Supported file types for processing"""
    MARKDOWN = ['.md', '.markdown']
    TEXT = ['.txt', '.text']
    WORD = ['.doc', '.docx']
    PDF = ['.pdf']
    ASCIIDOC = ['.adoc', '.adocx', '.asciidoc']


class FileProcessor:
    """Handles processing of various file types into plain text"""
    
    def __init__(self):
        self.supported_extensions = self._get_supported_extensions()
    
    def _get_supported_extensions(self) -> set:
        """Get all supported file extensions"""
        extensions = set()
        for file_type in SupportedFileTypes:
            extensions.update(file_type.value)
        return extensions
    
    def process_file(self, uploaded_file) -> str:
        """
        Process uploaded file and return its content as text
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: The content of the file as text
            
        Raises:
            ValueError: If file type is not supported
            Exception: If file processing fails
        """
        if uploaded_file is None:
            raise ValueError("No file provided")
        
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            # Create a bytes buffer from the uploaded file
            file_content = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer for potential re-reading
            
            # Process based on file type
            if file_extension in SupportedFileTypes.PDF.value:
                return self._process_pdf(file_content)
            elif file_extension in SupportedFileTypes.WORD.value:
                return self._process_word(file_content, file_extension)
            elif file_extension in SupportedFileTypes.MARKDOWN.value:
                return self._process_markdown(file_content)
            elif file_extension in SupportedFileTypes.TEXT.value:
                return self._process_text(file_content)
            elif file_extension in SupportedFileTypes.ASCIIDOC.value:
                return self._process_asciidoc(file_content)
            else:
                # Fallback to text processing
                return self._process_text(file_content)
        
        except Exception as e:
            raise Exception(f"Error processing file {uploaded_file.name}: {str(e)}")
    
    def _detect_encoding(self, content: bytes) -> str:
        """Detect text encoding of byte content"""
        if chardet:
            try:
                result = chardet.detect(content)
                return result['encoding'] or 'utf-8'
            except:
                return 'utf-8'
        else:
            # Fallback encoding detection
            try:
                content.decode('utf-8')
                return 'utf-8'
            except UnicodeDecodeError:
                return 'latin-1'
    
    def _process_text(self, content: bytes) -> str:
        """Process plain text files"""
        encoding = self._detect_encoding(content)
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            return content.decode('utf-8', errors='replace')
    
    def _process_markdown(self, content: bytes) -> str:
        """Process Markdown files - return as plain text for now"""
        # For now, just treat as text. Could convert to HTML later if needed
        text_content = self._process_text(content)
        
        # Optional: You could convert markdown to text here
        # For Jira tickets, raw markdown is probably fine
        return text_content
    
    def _process_pdf(self, content: bytes) -> str:
        """Process PDF files using PyPDF2"""
        if not PyPDF2:
            raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            if not text_content:
                raise Exception("No text content found in PDF")
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            raise Exception(f"Failed to process PDF: {str(e)}")
    
    def _process_word(self, content: bytes, extension: str) -> str:
        """Process Word documents (.doc and .docx)"""
        
        if extension == '.docx':
            if not docx:
                raise Exception("python-docx not installed. Install with: pip install python-docx")
            
            try:
                doc_file = io.BytesIO(content)
                doc = docx.Document(doc_file)
                
                text_content = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract table content
                for table in doc.tables:
                    table_content = []
                    for row in table.rows:
                        row_content = []
                        for cell in row.cells:
                            row_content.append(cell.text.strip())
                        if any(row_content):  # Only add non-empty rows
                            table_content.append(" | ".join(row_content))
                    
                    if table_content:
                        text_content.append("\n--- Table ---\n" + "\n".join(table_content))
                
                if not text_content:
                    raise Exception("No text content found in Word document")
                
                return "\n\n".join(text_content)
            
            except Exception as e:
                raise Exception(f"Failed to process .docx document: {str(e)}")
        
        elif extension == '.doc':
            # Try multiple approaches for .doc files
            return self._process_doc_file(content)
    
    def _process_doc_file(self, content: bytes) -> str:
        """Process .doc files using multiple fallback strategies"""
        
        # Strategy 1: Check if it's actually HTML content (common for Jira exports)
        try:
            text_content = content.decode('utf-8', errors='ignore')
            if '<html' in text_content.lower() or '<!doctype html' in text_content.lower():
                return self._extract_text_from_html(text_content)
        except:
            pass
        
        # Strategy 2: Try docx2txt if available
        if docx2txt:
            try:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                    tmp_file.write(content)
                    tmp_file.flush()
                    
                    text = docx2txt.process(tmp_file.name)
                    
                    # Clean up temp file
                    os.unlink(tmp_file.name)
                    
                    if text and text.strip():
                        return text
            except Exception as e:
                print(f"docx2txt failed: {e}")
        
        # Strategy 3: Try as plain text with various encodings
        encodings_to_try = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1', 'ascii']
        
        for encoding in encodings_to_try:
            try:
                text = content.decode(encoding, errors='ignore')
                # Remove HTML tags if present and extract readable text
                cleaned_text = self._clean_text_content(text)
                if len(cleaned_text.strip()) > 50:  # Must have substantial content
                    return cleaned_text
            except:
                continue
        
        # Strategy 4: Last resort - extract any readable text
        try:
            # Filter out non-printable characters and extract readable text
            readable_chars = ''.join(chr(b) for b in content if 32 <= b <= 126 or b in [9, 10, 13])
            if len(readable_chars.strip()) > 50:
                return readable_chars
        except:
            pass
        
        raise Exception("Could not extract text from .doc file. Please try converting to .docx format for better compatibility.")
    
    def _extract_text_from_html(self, html_content: str) -> str:
        """Extract text from HTML content"""
        
        if BeautifulSoup:
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text
                text = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text
            except:
                pass
        
        # Fallback: Simple HTML tag removal
        import re
        # Remove HTML tags
        text = re.sub('<.*?>', '', html_content)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _clean_text_content(self, text: str) -> str:
        """Clean and filter text content"""
        import re
        
        # Remove HTML tags if present
        text = re.sub('<.*?>', '', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if ord(char) >= 32 or char in ['\n', '\t'])
        
        # Clean up multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def _process_asciidoc(self, content: bytes) -> str:
        """Process AsciiDoc files"""
        # For now, treat as text since we want to preserve AsciiDoc syntax
        return self._process_text(content)
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get information about the uploaded file"""
        if uploaded_file is None:
            return {}
        
        file_size = len(uploaded_file.read())
        uploaded_file.seek(0)  # Reset file pointer
        
        return {
            'name': uploaded_file.name,
            'size': file_size,
            'type': uploaded_file.type,
            'extension': Path(uploaded_file.name).suffix.lower(),
            'size_mb': round(file_size / (1024 * 1024), 2)
        }
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file is supported"""
        extension = Path(filename).suffix.lower()
        return extension in self.supported_extensions
    
    def get_supported_extensions_list(self) -> list:
        """Get list of supported extensions for display"""
        return sorted(list(self.supported_extensions))


def test_file_processor():
    """Test function for file processor"""
    processor = FileProcessor()
    
    print("Supported extensions:", processor.get_supported_extensions_list())
    
    # Test with sample text
    sample_text = "This is a test document\n\nWith multiple paragraphs."
    sample_bytes = sample_text.encode('utf-8')
    
    result = processor._process_text(sample_bytes)
    print("Text processing test:", result[:50] + "..." if len(result) > 50 else result)


if __name__ == "__main__":
    test_file_processor()
